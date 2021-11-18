from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

# Profile picture
document.add_picture('john.jpg', width=Inches(2.0))

# Name, Phone number and Email details
name = input('Enter your name: ')
speak('Hello ' + name)

phone_number = input('Enter your phone number: ')
email = input('Enter your email: ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('About me')
about_me = input('Tell about yourself: ')
document.add_paragraph(about_me)

# Work experience
document.add_heading('Work experience')

p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('From date: ')
to_date = input('To date: ')

# Adding text to exsisting paragraph
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No: ')

    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company: ')
        from_date = input('From date: ')
        to_date = input('To date: ')

        # Adding text to exsisting paragraph
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ': ')
        p.add_run(experience_details)

    else:
        break

# Skills
document.add_heading('Skills')
skill = input('Enter skill: ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or NO: ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter skill: ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


document.save('cv.docx')